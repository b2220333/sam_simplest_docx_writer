#!/usr/bin/env python3
#encoding: utf-8

#開啟docx（聽說不能開doc檔，但實際開好像沒問題XD）
from docx import Document
doc1 = Document('../linux應用_v9.75_simple.doc')
#print(f'Doc={Doc}')

#列出所有大綱
#for paragraph in doc1.paragraphs:
    #第1階會輸出：Heading 1
    #第2階會輸出：Heading 2
    #...
    #依此類推
    #print(f'大綱階層為：{paragraph.style.name}')
    #print(f'內文為：{paragraph.text}')

#exit(0)
#建立Tkinter GUI
import tkinter
from tkinter import ttk
from tkinter import *

root = tkinter.Tk()
root.title('sam的Word處理器')
root.geometry("1200x800")  # Width x Height

#建立左側欄位
frame1 = Frame(root)
frame1.pack(side=LEFT)


#設定每列的高度
style_Treeview = ttk.Style(root)
style_Treeview.configure('Treeview', rowheight=40)
#設定欄位文字大小
style_Heading = ttk.Style()
style_Heading.configure("Treeview.Heading", font=(None, 15))



#千萬別加show='headings'，否則會縮小到看不到展開的+號
#tree = ttk.Treeview(root, columns=['1'], show='headings')
#tree = ttk.Treeview(root, columns=['欄位ID 1','欄位ID 2'], height=40)
tree = ttk.Treeview(root, height=40)
tree.column("#0",minwidth=0,width=500)
#anchor='center'表示置中對齊
#tree.column('1', width=100, anchor='center')
#tree.heading('1', text='大綱')
#tree["columns"]=("one","two")
#tree.column("欄位ID 1", width=100)
#tree.column("欄位ID 2", width=100)
#tree.heading("欄位ID 1", text="欄位文字1")
#tree.heading("欄位ID 2", text="欄位文字2")

"""
#第一階層-1
#第二欄位表順序（0表最優先）
tree.insert("" , 0,'階層1-1 ID', text="階層1-1文字", values=['a','b'])

#第一階層-2
apple = tree.insert("", 2, "階層1-2 ID", text="階層1-2文字")
#第二階層
tree.insert(apple, 0,"階層2-1 ID1", text="階層2-1文字a", values=['c','d'])
tree.insert(apple, 1,"階層2-1 ID2", text="階層2-1文字b", values=['e','f'])

from PIL import Image, ImageTk

linux_img2 = Image.open('/home/sam/PycharmProjects/linux.png')
linux_img2 = linux_img2.resize((30,30), Image.ANTIALIAS)
linux_img2 = ImageTk.PhotoImage(linux_img2)
tree.insert(apple, 3,"階層2-1 ID2 pic2", text='linux.png', open=True, image=linux_img2,
                 value=['0','1'])
#第一階層-3
tree.insert('',3,"階層1-3 ID", text="階層1-3文字", values=['g','h'])

"""

#準備建立獨特的ID
ID = 0

#若N不變表示多了同階
#若N變大表示降了一階（如階層一->階層二）
#若N變小表示升了一階（如階層二->階層一）
#取得所有大綱
#先建立特殊的首個第一階層
previous_level = 0
stack_var_level = []

#內文
g_body_text = []


#此處是用來parse 底層docx的xml來抓出圖片用
g_image = []
g_image_index = 0
import re
from os.path import basename
pattern = re.compile('rId\d+')

#此處是用來抓出超連結
from docx.opc.constants import RELATIONSHIP_TYPE as RT
g_hyperlink_index = 0
#取得所有超連結
all_hyperlinks=[]
rels = doc1.part.rels
for rel in rels:
    print(f'rel={rel}')
    if rels[rel].reltype == RT.HYPERLINK:
        print("\n 超連結的id為", rel, "，超連結為", rels[rel]._target)
        all_hyperlinks.append(rels[rel]._target)

for paragraph in doc1.paragraphs:
    #第1階會輸出：Heading 1
    #第2階會輸出：Heading 2
    #...
    #依此類推
    print('-----------------------------------------')
    #print(f'大綱階層為：{paragraph.style.name}')
    #print(f'階層標題為：{paragraph.text}')
    #print(f'stack_var_level={stack_var_level}')
    #print(f'previous_level={previous_level}')

    #開始取得Heading N的數字N
    #str[str.index(' ')+1:]
    if previous_level == 0:
        #表示目前大綱尚未建立任何內容
        #var_level = tree.insert("", ID, f'{ID}', text={paragraph.text}, values=[paragraph.text, 'b'])
        var_level = tree.insert("", ID, f'{ID}', text={paragraph.text})
        stack_var_level.append(var_level)
        print(f'stack_var_level={stack_var_level}')
        previous_level = 1
        ID = 1
        # 每抓到一個paragraph就要在g_body_text多新增一個內容
        # 才不會要儲存某個ID的內文時發現沒有空間可存
        g_body_text.append('')
        print(f'len(g_body_text)={len(g_body_text)}')
    else:
        #表示目前至少有一個大綱
        #但若得到的資訊是Normal則此奇怪欄位須跳過，因為也不會有任何內容
        #若是Body Text也屬於內文範疇，所以在大綱方面應忽略
        if paragraph.style.name == 'Normal' or paragraph.style.name == 'Body Text':
            if paragraph.text!= '':
                #不是空字串就可以正常顯示出來
                print(f'階層內文為：{paragraph.text}')
                #因為一定有大綱才有內文，因此也一定先建立了g_body_text
                #所以只須將body的文字全部存入到g_body_text[-1]中即可。
                g_body_text[-1] = g_body_text[-1] + paragraph.text + '\n'
                print(f'g_body_text[-1]={g_body_text[-1]}')
            else:
                #若是空字串表示抓到的是不一定是圖，也有可能是超連結（目前程式不處理非文字和圖和超連結之外的結構如表格）
                local_image_list1 = list()
                #實做方式很複雜，必須深入xml層去抓影像資料才行（所以須先透過run來偵測）
                for run in paragraph.runs:
                    if run.text != '':
                        print(f'run.text={run.text}')
                        local_image_list1.append(run.text)
                    else:
                        print(f'run.text={run.text}')
                        # b.append(pattern.search(run.element.xml))
                        contentID = pattern.search(run.element.xml).group(0)
                        try:
                            contentType = doc1.part.related_parts[contentID].content_type
                            print(f'contentType={contentType}')
                        except KeyError as e:
                            print(e)
                            continue
                        if not contentType.startswith('image'):
                            continue
                        imgName = basename(doc1.part.related_parts[contentID].partname)
                        imgData = doc1.part.related_parts[contentID].blob
                        local_image_list1.append(imgData)
                if local_image_list1!=[]:
                    #不是空表示有抓到圖
                    print(f'準備顯示第{g_image_index}張圖！')
                    #for image in doc1.inline_shapes:
                        #print(f'影像寬度為{image.width},影像高度為{image.height}')
                    print(f'影像寬度為{doc1.inline_shapes[g_image_index].width}')
                    print(f'影像高度為{doc1.inline_shapes[g_image_index].height}')
                    g_image_index = g_image_index+1
                else:
                    #此時發現是空，表示此處沒有圖
                    print('***********此處既不是文字也非圖，那就假設是超連結吧！**************')
                    #試著印出超連結
                    print(f'超連結為{all_hyperlinks[g_hyperlink_index]}')
                    g_hyperlink_index += 1
                    #目前假設超連結也屬於內文的一部分，因此作法比照內文方式處理
                    #因為一定有大綱才有內文，因此也一定先建立了g_body_text
                    #所以只須將body的文字全部存入到g_body_text[-1]中即可。
                    g_body_text[-1] = g_body_text[-1] + paragraph.text + '\n'
                    print(f'g_body_text[-1]={g_body_text[-1]}')
                    #pass
            continue
        print(f'大綱階層為：{paragraph.style.name}')
        print(f'階層標題為：{paragraph.text}')
        # print(f'stack_var_level={stack_var_level}')
        print(f'previous_level={previous_level}')

        now_level = int(paragraph.style.name[paragraph.style.name.index(' ')+1:])
        print(f'now_level={now_level}')
#        if now_level == 2:
#            print(f'大綱階層為：{paragraph.style.name}')
#            print(f'內文為：{paragraph.text}')
#            print(f'stack_var_level={stack_var_level}')

        if now_level > previous_level:
            if now_level-previous_level != 1:
                print('請修補文件為正確大綱階層（一次只能升一階或降一階）')
                exit(-1)
            #準備降階
            #tree.insert("", 0, '階層1-1 ID', text="階層1-1文字", values=['a', 'b'])
            #var_level = tree.insert(var_level, ID, f'{ID}', text=paragraph.style.name+':'+paragraph.text, values=[paragraph.text, 'b'])
            var_level = tree.insert(var_level, ID, f'{ID}', text=paragraph.style.name + ':' + paragraph.text)
            stack_var_level.append(var_level)
            print(f'stack_var_level={stack_var_level}')
            previous_level = now_level
            ID = ID + 1
            # 每抓到一個paragraph就要在g_body_text多新增一個內容
            # 才不會要儲存某個ID的內文時發現沒有空間可存
            g_body_text.append('')
            print(f'len(g_body_text)={len(g_body_text)}')

        elif now_level == previous_level:
            #因為目前抓到同階層的，所以需要把前面同階層的變數丟棄
            #這樣才能抓到上一階層的來降階，才能夠產生正確的大綱
            print('準備要pop')
            stack_var_level.pop()
            #var_level = tree.insert(stack_var_level[-1], ID, f'{ID}', text=paragraph.style.name+':'+paragraph.text, values=[paragraph.text, 'b'])
            var_level = tree.insert(stack_var_level[-1], ID, f'{ID}', text=paragraph.style.name + ':' + paragraph.text)
            #最後因為可能此階層也會有下一階，因此當然變數仍要push
            stack_var_level.append(var_level)
            print(f'stack_var_level={stack_var_level}')
            previous_level = now_level
            ID = ID + 1
            # 每抓到一個paragraph就要在g_body_text多新增一個內容
            # 才不會要儲存某個ID的內文時發現沒有空間可存
            g_body_text.append('')
            print(f'len(g_body_text)={len(g_body_text)}')

        elif now_level < previous_level:
            if previous_level-now_level != 1:
                print('請修補文件為正確大綱階層（一次只能升一階或降一階）')
                exit(-1)
            #準備升階
            #pop一次可以建立同階層
            print('準備要pop')
            stack_var_level.pop()
            #<1要多pop共1次
            #<2要多pop共2次
            for _ in range(previous_level-now_level):
                print('準備要pop')
                stack_var_level.pop()
            #print(f'stack_var_level={stack_var_level}')
            #var_level = tree.insert(stack_var_level[-1], ID, f'{ID}', text=paragraph.style.name+':'+paragraph.text, values=[paragraph.text, 'b'])
            var_level = tree.insert(stack_var_level[-1], ID, f'{ID}', text=paragraph.style.name + ':' + paragraph.text)
            #最後因為可能此階層也會有下一階，因此當然變數仍要push
            stack_var_level.append(var_level)
            print(f'stack_var_level={stack_var_level}')
            previous_level = now_level
            ID = ID + 1
            # 每抓到一個paragraph就要在g_body_text多新增一個內容
            # 才不會要儲存某個ID的內文時發現沒有空間可存
            g_body_text.append('')
            print(f'len(g_body_text)={len(g_body_text)}')

        #break
#tree.grid()
#tree.pack()
tree.pack(side=LEFT)



#######################################################################################
#建立右側欄位
frame2 = Frame(root)
frame2.pack(side=RIGHT)

text = Text(root, width=80)

#讓text靠左
text.pack(side=LEFT, fill=Y, expand=True)
scrollbar = Scrollbar(root, orient="vertical")
scrollbar.config(command=text.yview)
#讓捲軸靠右
scrollbar.pack(side=RIGHT, fill=Y, expand=True)

################################事件處理############################

def treeviewClick(event):  # 單擊
    now_ID = int(tree.selection()[0])
    print(f'單擊ID{now_ID}')
    print(f'該大綱的內文為：{g_body_text[now_ID]}')
    #text.insert('insert', 'Hello sam!!')
    #先清空文字框
    text.delete('1.0', END)
    #再加入目前大綱中的文字
    text.insert('insert', g_body_text[now_ID])
    #for item in tree.selection():
    #    item_text = tree.item(item, "values")
    #    print(item_text[0])  # 輸出所選行的第一列的值


tree.bind('<ButtonRelease-1>', treeviewClick)  # 綁定單擊離開事件===========

root.mainloop()
